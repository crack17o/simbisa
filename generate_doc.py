"""
Génère le document de présentation Simbisa (non-technique).
Usage : python generate_doc.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Palette ──────────────────────────────────────────────────────────────────
BLEU_RAWBANK   = RGBColor(0x00, 0x38, 0x7A)   # bleu foncé institution
BLEU_CLAIR     = RGBColor(0x00, 0x7A, 0xC1)   # accent
GRIS_FOND      = RGBColor(0xF4, 0xF6, 0xF9)
GRIS_TEXTE     = RGBColor(0x44, 0x44, 0x55)
BLANC          = RGBColor(0xFF, 0xFF, 0xFF)
VERT_SCORE     = RGBColor(0x1A, 0x8A, 0x5A)
ORANGE_ALERTE  = RGBColor(0xE0, 0x7A, 0x00)


# ─── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if side in kwargs:
            tag = OxmlElement(f'w:{side}')
            tag.set(qn('w:val'),   kwargs[side].get('val', 'single'))
            tag.set(qn('w:sz'),    kwargs[side].get('sz', '4'))
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), kwargs[side].get('color', '000000'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def no_space_para(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'),  '0')
    pPr.append(spacing)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '007AC1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLEU_RAWBANK
    add_horizontal_rule(doc)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLEU_CLAIR
    return p


def body(doc, text, color=None):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        if color:
            run.font.color.rgb = color
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def info_box(doc, title, text, color_hex='E8F0FA', title_color=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color_hex)
    cell.width = Inches(6.3)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p0 = cell.paragraphs[0]
    no_space_para(p0)
    r0 = p0.add_run(f'  {title}')
    r0.bold = True
    r0.font.size = Pt(11)
    r0.font.color.rgb = title_color or BLEU_RAWBANK

    p1 = cell.add_paragraph(f'  {text}')
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after  = Pt(6)
    for run in p1.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = GRIS_TEXTE

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def score_bar_table(doc, label, value, max_val, color_hex):
    """Ligne de score visuelle : label | barre | valeur."""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    c0, c1, c2 = tbl.row_cells(0)
    c0.width = Inches(1.8)
    c1.width = Inches(3.5)
    c2.width = Inches(1.0)

    # Label
    c0.paragraphs[0].add_run(label).font.size = Pt(10)
    c0.paragraphs[0].paragraph_format.space_after = Pt(2)

    # Barre de fond (gris clair)
    set_cell_bg(c1, 'E8ECF0')
    # Simulé par texte
    filled = int((value / max_val) * 30)
    bar = '█' * filled + '░' * (30 - filled)
    r = c1.paragraphs[0].add_run(bar)
    r.font.size  = Pt(7)
    r.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16),
    )
    c1.paragraphs[0].paragraph_format.space_after = Pt(2)

    # Valeur
    r2 = c2.paragraphs[0].add_run(f'{value}/{max_val}')
    r2.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16),
    )
    c2.paragraphs[0].paragraph_format.space_after = Pt(2)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def cover_page(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(40)

    # Logo texte / titre institution
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_inst.add_run('RAWBANK  ·  SIMBISA')
    r.bold      = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLEU_CLAIR
    r.font.all_caps  = True

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Titre principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_title.add_run('SIMBISA')
    r1.bold = True
    r1.font.size = Pt(40)
    r1.font.color.rgb = BLEU_RAWBANK

    # Sous-titre
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run('Système Intelligent de Scoring et de Micro-crédit pour Banques en Afrique')
    r2.font.size = Pt(15)
    r2.font.color.rgb = BLEU_CLAIR
    r2.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(30)

    # Ligne décorative
    add_horizontal_rule(doc)

    # Accroche
    p_hook = doc.add_paragraph()
    p_hook.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_hook.add_run(
        'Une plateforme numérique complète pour accélérer l\'accès au micro-crédit\n'
        'par l\'intelligence artificielle — Mobile · Web · USSD'
    )
    r3.font.size = Pt(12)
    r3.font.color.rgb = GRIS_TEXTE
    r3.italic = True
    p_hook.paragraph_format.space_before = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(30)

    # Métadonnées
    meta = [
        ('Projet',      'TFC / Implémentation – Rawbank RDC'),
        ('Version',     '1.0 — Juillet 2026'),
        ('Auteur',      'Joël Stone'),
        ('Public cible','Direction, partenaires, membres du jury (non-techniques)'),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        row = tbl.rows[i]
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.3)
        set_cell_bg(row.cells[0], 'E8F0FA')
        r_k = row.cells[0].paragraphs[0].add_run(k)
        r_k.bold = True
        r_k.font.size = Pt(10)
        r_k.font.color.rgb = BLEU_RAWBANK
        r_v = row.cells[1].paragraphs[0].add_run(v)
        r_v.font.size = Pt(10)
        r_v.font.color.rgb = GRIS_TEXTE

    doc.add_page_break()


def section_contexte(doc):
    heading1(doc, '1. Contexte du projet')

    body(doc,
        'La République Démocratique du Congo abrite une économie majoritairement informelle : '
        'plus de 80 % de la population active travaille sans contrat de travail formel, '
        'sans fiches de paie ni relevés bancaires traditionnels. Pourtant, ces mêmes '
        'personnes réalisent chaque jour des millions de transactions sur les réseaux '
        'Mobile Money — M-Pesa, Orange Money, Airtel Money, Africell.')

    body(doc,
        'Rawbank, première banque privée du pays, a lancé Simbisa pour répondre à une '
        'réalité simple : les outils bancaires classiques excluent ceux qui en ont le plus '
        'besoin. Un commerçant de Gombe, une enseignante de Lemba ou un artisan de Makala '
        'ont souvent un revenu réel et régulier — mais aucune preuve que les banques '
        'savent lire.')

    heading2(doc, 'Qu\'est-ce que Simbisa ?')
    body(doc,
        'Simbisa est une plateforme numérique qui permet à Rawbank d\'évaluer, en quelques '
        'minutes, la capacité de remboursement d\'un client à partir de données réelles : '
        'son historique Mobile Money, son comportement d\'épargne, et son profil. '
        'Un score de crédit est calculé automatiquement par l\'intelligence artificielle, '
        'et une décision est rendue — avec une explication claire.')

    info_box(doc,
        'En résumé',
        'Simbisa transforme l\'historique Mobile Money d\'un client en une évaluation de '
        'crédit fiable, réduisant le délai d\'analyse de plusieurs jours à quelques minutes.',
        'E8F4EE', VERT_SCORE)

    bullet(doc, 'Interface mobile pour les clients (Android / iOS)')
    bullet(doc, 'Tableau de bord web pour les équipes Rawbank')
    bullet(doc, 'Canal USSD (*123#) pour les zones à faible connectivité')
    bullet(doc, 'Moteur de scoring IA entraîné sur les données mobiles congolaises')


def section_problematique(doc):
    heading1(doc, '2. Problématique')

    body(doc,
        'Avant Simbisa, l\'octroi d\'un micro-crédit chez Rawbank suivait un processus '
        'entièrement manuel : réception du dossier papier, vérification des pièces, '
        'analyse du conseiller, validation du responsable. Ce processus prenait en moyenne '
        '3 à 5 jours ouvrables.')

    heading2(doc, 'Les quatre problèmes majeurs identifiés')

    problems = [
        ('Délais excessifs',
         'Un commerçant qui a besoin de liquidités pour réapprovisionner son stock '
         'ne peut pas attendre 5 jours. Le crédit informel — coûteux et risqué — '
         'reste souvent son seul recours.'),
        ('Biais humains',
         'Sans critères objectifs et standardisés, deux agents peuvent rendre des '
         'décisions différentes pour des profils identiques. La subjectivité nuit '
         'à l\'équité et expose Rawbank à des risques de plainte.'),
        ('Données Mobile Money ignorées',
         'Un client peut avoir 2 ans d\'historique régulier sur M-Pesa — réceptions '
         'de salaire, paiements de factures, épargne — sans que ces données ne soient '
         'jamais lues lors de l\'analyse de crédit.'),
        ('Traçabilité insuffisante',
         'Sans journal numérique des décisions, l\'audit interne et la conformité '
         'réglementaire (BCC) deviennent difficiles à garantir.'),
    ]

    tbl = doc.add_table(rows=len(problems), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (titre, desc) in enumerate(problems):
        row = tbl.rows[i]
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.8)
        set_cell_bg(row.cells[0], 'FFF3E0')
        r0 = row.cells[0].paragraphs[0].add_run(titre)
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = ORANGE_ALERTE
        row.cells[0].paragraphs[0].paragraph_format.space_after = Pt(3)
        r1 = row.cells[1].paragraphs[0].add_run(desc)
        r1.font.size = Pt(10)
        r1.font.color.rgb = GRIS_TEXTE
        row.cells[1].paragraphs[0].paragraph_format.space_after = Pt(3)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def section_situation(doc):
    heading1(doc, '3. Situation actuelle avec Simbisa')

    body(doc,
        'Simbisa est aujourd\'hui une plateforme opérationnelle, déployée sur le serveur '
        'de Rawbank. Elle couvre l\'ensemble du cycle de vie du micro-crédit : de la '
        'demande du client jusqu\'à la clôture du crédit, en passant par le scoring IA, '
        'la décision humaine et le suivi des remboursements.')

    phases = [
        ('Phase 1 : Inscription & KYC',
         'Le client s\'inscrit via l\'application mobile. Il renseigne ses informations '
         'personnelles et soumet une pièce d\'identité (carte d\'électeur, passeport ou permis '
         'de conduire). L\'agent de crédit valide le dossier en agence.'),
        ('Phase 2 : Demande de crédit',
         'Le client soumet une demande depuis l\'app mobile ou via USSD (*123#). '
         'Il indique le montant souhaité, la durée et le motif. La plateforme vérifie '
         'immédiatement les prérequis (KYC valide, âge 20-60 ans, pas de crédit en cours).'),
        ('Phase 3 : Scoring automatique',
         'En quelques secondes, quatre moteurs analysent le profil du client et produisent '
         'un score de 0 à 100. Une décision préliminaire est générée automatiquement, '
         'accompagnée d\'une explication rédigée par l\'intelligence artificielle.'),
        ('Phase 4 : Décision & déblocage',
         'Selon le montant demandé, la décision est soit automatique, soit validée par un '
         'agent ou un responsable crédit. Une fois approuvé, le montant est crédité sur le '
         'wallet du client.'),
        ('Phase 5 : Suivi & remboursement',
         'Le client rembourse selon l\'échéancier, via Mobile Money ou en agence. '
         'L\'application affiche les mensualités dues, les paiements effectués et '
         'le solde restant.'),
    ]

    for i, (phase, desc) in enumerate(phases):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.3)
        r_num = p.add_run(f'{i + 1}  ')
        r_num.bold = True
        r_num.font.size = Pt(13)
        r_num.font.color.rgb = BLEU_CLAIR
        r_title = p.add_run(phase)
        r_title.bold = True
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = BLEU_RAWBANK

        p2 = doc.add_paragraph(desc)
        p2.paragraph_format.left_indent  = Cm(1.0)
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after  = Pt(8)
        for run in p2.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = GRIS_TEXTE


def section_utilisateurs(doc):
    heading1(doc, '4. Actions de chaque utilisateur')

    body(doc,
        'La plateforme distingue six rôles aux responsabilités clairement délimitées. '
        'Chaque utilisateur dispose d\'un accès personnalisé — application mobile pour '
        'les clients, tableau de bord web pour les équipes Rawbank.')

    roles = [
        (
            'Client',
            'Application mobile (Android/iOS) + USSD *123#',
            BLEU_CLAIR,
            'E8F0FA',
            [
                'S\'inscrire et soumettre ses pièces d\'identité pour le KYC',
                'Consulter son score de crédit en temps réel (anneau de score)',
                'Soumettre une demande de micro-crédit (montant, durée, motif)',
                'Suivre l\'avancement de sa demande et la décision rendue',
                'Gérer son épargne : créer un objectif, déposer, suivre la progression',
                'Consulter son solde wallet (USD/CDF) et son historique',
                'Rembourser ses mensualités via Mobile Money',
                'Accéder aux services via USSD depuis n\'importe quel téléphone',
            ],
        ),
        (
            'Agent de crédit',
            'Tableau de bord web — gestion des dossiers clients',
            RGBColor(0x1A, 0x6B, 0x3A),
            'EBF5EE',
            [
                'Valider les dossiers KYC des clients de sa commune',
                'Consulter les demandes de crédit en attente d\'analyse',
                'Lire le score IA et l\'explication générée pour chaque demande',
                'Prendre une décision manuelle (approuver / rejeter) avec motif',
                'Gérer les exceptions : signaler une anomalie ou escalader au responsable',
                'Suivre les crédits en cours et les remboursements de ses clients',
            ],
        ),
        (
            'Responsable crédit',
            'Tableau de bord web — supervision et validation haute valeur',
            RGBColor(0x6B, 0x38, 0x00),
            'FFF8EE',
            [
                'Valider les demandes dépassant le plafond agent (> 400 USD)',
                'Examiner et trancher les exceptions remontées par les agents',
                'Consulter les statistiques globales : taux d\'approbation, montants engagés',
                'Superviser les performances de l\'équipe d\'agents de crédit',
            ],
        ),
        (
            'Analyste risque',
            'Tableau de bord web — pilotage du modèle IA',
            RGBColor(0x55, 0x00, 0x8A),
            'F5EEFF',
            [
                'Surveiller les performances du modèle XGBoost (précision, rappel, F1)',
                'Consulter les attributions SHAP (quelles variables influencent le score)',
                'Analyser la distribution des scores et les taux de défaut',
                'Relancer manuellement un ré-entraînement du modèle sur les nouvelles données',
                'Ajuster les paramètres de scoring si nécessaire',
            ],
        ),
        (
            'Auditeur',
            'Tableau de bord web — traçabilité et conformité',
            RGBColor(0x6B, 0x20, 0x20),
            'FFF0EE',
            [
                'Consulter l\'historique complet de toutes les décisions de crédit',
                'Accéder au détail de chaque décision : client, score, moteur IA, explication',
                'Vérifier la cohérence entre recommandation IA et décision humaine',
                'Exporter les journaux pour les rapports de conformité BCC',
                'Détecter les anomalies ou décisions suspectes',
            ],
        ),
        (
            'Administrateur',
            'Tableau de bord web — configuration de la plateforme',
            RGBColor(0x20, 0x20, 0x20),
            'F4F4F4',
            [
                'Gérer les comptes utilisateurs et les rôles',
                'Configurer les plafonds de crédit par niveau de compte',
                'Mettre à jour le taux de change USD/CDF',
                'Superviser la santé générale de la plateforme',
            ],
        ),
    ]

    for role_name, interface, title_color, bg_hex, actions in roles:
        # En-tête du rôle
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_bg(cell, bg_hex)
        p0 = cell.paragraphs[0]
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after  = Pt(0)
        r0 = p0.add_run(f'  {role_name}')
        r0.bold = True
        r0.font.size = Pt(12)
        r0.font.color.rgb = title_color
        p1 = cell.add_paragraph(f'  {interface}')
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after  = Pt(4)
        for run in p1.runs:
            run.italic = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = GRIS_TEXTE

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        for action in actions:
            bullet(doc, action)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)


def section_score(doc):
    heading1(doc, '5. Fonctionnement du score de crédit')

    body(doc,
        'Le score Simbisa est une note de 0 à 100 qui mesure la probabilité qu\'un client '
        'rembourse son crédit sans incident. Il n\'est pas calculé par un humain : quatre '
        'moteurs indépendants analysent des aspects différents du profil, et leurs '
        'résultats sont combinés pour produire un score final.')

    # Les 4 moteurs
    moteurs = [
        ('Moteur 1 — Règles & KYC', '1A6B3A',
         25, 25,
         'Vérifie les conditions préalables obligatoires.',
         [
             'KYC valide (pièce d\'identité vérifiée par un agent)',
             'Âge compris entre 20 et 60 ans (conformité BCC)',
             'Aucun crédit en cours sur la même devise',
             'Ancienneté sur la plateforme ≥ 30 jours',
             'Montant demandé conforme au plafond du niveau de compte',
         ]),
        ('Moteur 2 — Mobile Money', '007AC1',
         30, 30,
         'Analyse l\'historique Mobile Money des 90 derniers jours.',
         [
             'Volume total des entrées et des sorties',
             'Ratio entrées/sorties (une personne qui reçoit plus qu\'elle ne dépense)',
             'Régularité des transactions (fréquence mensuelle stable)',
             'Présence de dépôts récurrents (revenu régulier)',
             'Absence de comportements à risque (retraits massifs, solde négatif)',
         ]),
        ('Moteur 3 — Comportemental', 'E07A00',
         20, 20,
         'Évalue le comportement du client sur la plateforme Simbisa.',
         [
             'Respect des précédents remboursements (si crédit antérieur)',
             'Régularité des versements sur le compte épargne',
             'Progression vers les objectifs d\'épargne définis',
             'Ancienneté et fidélité sur la plateforme',
         ]),
        ('Moteur 4 — IA / XGBoost', '550088',
         25, 25,
         'Modèle d\'apprentissage automatique entraîné sur les décisions passées.',
         [
             'Calcule une probabilité de défaut (0 % à 100 %)',
             'Identifie les variables les plus déterminantes pour ce client spécifique',
             'Génère des attributions SHAP : quel facteur a le plus influencé la note',
             'Classe le risque : Faible / Moyen / Élevé / Très élevé',
         ]),
    ]

    for nom, color_hex, score, max_s, desc, items in moteurs:
        heading2(doc, nom)
        score_bar_table(doc, 'Poids max', score, 100, color_hex)
        body(doc, desc)
        for item in items:
            bullet(doc, item)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    heading2(doc, 'Score final et décision automatique')
    body(doc,
        'Le score final est la somme pondérée des quatre moteurs. Selon ce score, '
        'la plateforme génère une décision préliminaire :')

    decisions = [
        ('≥ 70 / 100', 'APPROUVÉ automatiquement', 'VERT_SCORE', '1A8A5A', 'D4EDDA'),
        ('50 – 69',    'À examiner par l\'agent', 'orange', 'E07A00', 'FFF3E0'),
        ('< 50',       'REFUSÉ automatiquement', 'rouge', 'CC2222', 'FFEEEE'),
    ]

    tbl = doc.add_table(rows=1 + len(decisions), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    headers = ['Score', 'Décision', 'Action requise']
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, '003C7A')
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = BLANC
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)

    rows_data = [
        ('≥ 70 / 100', 'Approuvé automatiquement',   'Aucune — crédit débloqué', 'D4EDDA', '1A8A5A'),
        ('50 – 69',    'Analyse requise',             'Agent / Responsable examine', 'FFF3E0', 'E07A00'),
        ('< 50',       'Refusé automatiquement',      'Motif communiqué au client',  'FFEEEE', 'CC2222'),
    ]
    for i, (score_txt, decision_txt, action_txt, bg, fg) in enumerate(rows_data):
        row = tbl.rows[i + 1]
        for j, txt in enumerate([score_txt, decision_txt, action_txt]):
            cell = row.cells[j]
            set_cell_bg(cell, bg if j == 1 else 'F8F8F8')
            r = cell.paragraphs[0].add_run(txt)
            r.font.size = Pt(10)
            if j == 1:
                r.bold = True
                r.font.color.rgb = RGBColor(int(fg[:2], 16), int(fg[2:4], 16), int(fg[4:], 16))
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    heading2(doc, 'Explication générée par l\'IA')
    info_box(doc,
        'Mémo IA — exemple',
        'Le profil de Jean Mutombo présente un score de 74/100. Son historique Mobile Money '
        'M-Pesa sur 90 jours montre des entrées régulières d\'environ 350 000 CDF/mois, '
        'un ratio entrées/sorties de 1,4 et aucune irrégularité majeure. Son niveau Pro '
        'autorise un crédit USD jusqu\'à 700 USD. Le montant demandé de 250 USD est '
        'conforme aux plafonds. Recommandation : APPROBATION.',
        'EEF4FF')


def section_identifiants(doc):
    heading1(doc, '6. Identifiants créés dans le système (données de démonstration)')

    body(doc,
        'Le système est livré avec des comptes de démonstration qui permettent de tester '
        'toutes les fonctionnalités immédiatement après le déploiement. '
        'Ces comptes sont créés par la commande : python manage.py seed_demo')

    info_box(doc,
        'Mot de passe commun (tous les comptes)',
        'Test123!     —     à changer impérativement en production',
        'FFF3E0', ORANGE_ALERTE)

    heading2(doc, 'Comptes du personnel Rawbank')

    staff_data = [
        ('Administrateur',       '+243 900 000 000', 'Admin Système',           'Toutes les fonctionnalités'),
        ('Agent crédit (Gombe)', '+243 900 000 002', 'Agent Kabongo',            'Gestion dossiers, KYC, décisions'),
        ('Agent crédit (Limete)','+243 900 000 006', 'Grace Limete Mputu',       'Gestion dossiers, KYC, décisions'),
        ('Agent crédit (Bandal)','+243 900 000 007', 'Jonas Bandal Mukasa',      'Gestion dossiers, KYC, décisions'),
        ('Responsable crédit',   '+243 900 000 003', 'Responsable Mukendi',      'Exceptions, validation haute valeur'),
        ('Analyste risque',      '+243 900 000 004', 'Analyste Tshilombo',       'Modèle IA, statistiques risque'),
        ('Auditeur',             '+243 900 000 005', 'Auditeur Ilunga',          'Journal décisions, conformité'),
    ]

    tbl = doc.add_table(rows=1 + len(staff_data), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-têtes
    for j, h in enumerate(['Rôle', 'Téléphone (login)', 'Nom complet', 'Accès']):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, '003C7A')
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = BLANC
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)

    for i, (role, tel, nom, acces) in enumerate(staff_data):
        row = tbl.rows[i + 1]
        bg = 'F0F4FA' if i % 2 == 0 else 'FFFFFF'
        for j, txt in enumerate([role, tel, nom, acces]):
            cell = row.cells[j]
            set_cell_bg(cell, bg)
            r = cell.paragraphs[0].add_run(txt)
            r.font.size = Pt(9.5)
            if j == 0:
                r.bold = True
                r.font.color.rgb = BLEU_RAWBANK
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    heading2(doc, 'Comptes clients de démonstration (20 profils)')

    client_data = [
        ('+243 810 000 001', 'Jean Kabila Mutombo',    'Gombe',     'Pro',      'M-Pesa',       '500 USD',  'Oui'),
        ('+243 820 000 002', 'Marie Mbuyi Lukusa',     'Limete',    'Standard', '—',            '0',        'Non'),
        ('+243 830 000 003', 'Paul Tshilombo Nkusu',   'Ngaliema',  'Pro+',     'Orange Money', '800 USD',  'Oui'),
        ('+243 840 000 004', 'Alice Ilunga Kasongo',   'Kinshasa',  'Premium',  'Airtel Money', '1 200 USD','Oui'),
        ('+243 850 000 005', 'Robert Kabongo Matamba', 'Gombe',     'Pro',      'Orange Money', '350 USD',  'Oui'),
        ('+243 860 000 006', 'Grace Mputu Nsimba',     'Kalamu',    'Standard', '—',            '200 USD',  'Non'),
        ('+243 870 000 007', 'Pierre Mukendi Kazadi',  'Lemba',     'Pro',      'Vodacom',      '650 USD',  'Oui'),
        ('+243 880 000 008', 'Sophie Ngoy Kabeya',     'Makala',    'Pro+',     'Africell',     '900 USD',  'Oui'),
        ('+243 890 000 009', 'David Luhaka Mwamba',    'Masina',    'Standard', 'Africell',     '0',        'Non'),
        ('+243 970 000 010', 'Claire Basila Tshianga', 'Gombe',     'Pro',      'Airtel Money', '700 USD',  'Oui'),
        ('+243 980 000 011', 'Michel Kahindo Bwana',   'Limete',    'Premium',  'Airtel Money', '1 500 USD','Oui'),
        ('+243 990 000 012', 'Jeanne Mbala Kilanda',   'Ngaliema',  'Standard', 'Airtel Money', '300 USD',  'Oui'),
        ('+243 810 000 013', 'Thomas Nzuzi Lutumba',   'Kinshasa',  'Pro',      'M-Pesa',       '450 USD',  'Oui'),
        ('+243 820 000 014', 'Anne Kasumba Mwana',     'Kasa-Vubu', 'Pro',      'Orange Money', '600 USD',  'Oui'),
        ('+243 830 000 015', 'Justin Malonda Mpaka',   'Gombe',     'Premium',  'Orange Money', '2 000 USD','Oui'),
        ('+243 840 000 016', 'Sandra Bikeka Luzolo',   'Lemba',     'Standard', 'Airtel Money', '0',        'Non'),
        ('+243 850 000 017', 'Eric Tshimanga Diakiese','Makala',    'Pro',      'Orange Money', '550 USD',  'Oui'),
        ('+243 860 000 018', 'Valérie Mabiku Nzinga',  'Masina',    'Pro',      '—',            '400 USD',  'Oui'),
        ('+243 870 000 019', 'Albert Kanku Mukeba',    'Limete',    'Pro+',     'Vodacom',      '1 100 USD','Oui'),
        ('+243 900 000 020', 'Céline Mwangu Tshika',   'Gombe',     'Pro',      'Africell',     '750 USD',  'Oui'),
    ]

    headers_c = ['Téléphone', 'Nom complet', 'Commune', 'Niveau', 'Opérateur MM', 'Revenu', 'KYC']
    tbl2 = doc.add_table(rows=1 + len(client_data), cols=len(headers_c))
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers_c):
        cell = tbl2.cell(0, j)
        set_cell_bg(cell, '003C7A')
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = BLANC
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)

    for i, row_data in enumerate(client_data):
        row = tbl2.rows[i + 1]
        bg = 'F0F4FA' if i % 2 == 0 else 'FFFFFF'
        for j, txt in enumerate(row_data):
            cell = row.cells[j]
            set_cell_bg(cell, bg if j != 6 else ('D4EDDA' if txt == 'Oui' else 'FFEEEE'))
            r = cell.paragraphs[0].add_run(txt)
            r.font.size = Pt(8.5)
            if j == 6:
                r.bold = True
                r.font.color.rgb = VERT_SCORE if txt == 'Oui' else RGBColor(0xCC, 0x22, 0x22)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    heading2(doc, 'PIN USSD des clients démo')
    info_box(doc,
        'PIN USSD commun',
        '0000   —   défini automatiquement par seed_demo pour les 20 clients. '
        'Chaque client peut le modifier lors de sa première connexion USSD réelle.',
        'E8F0FA')

    heading2(doc, 'Niveaux de compte et plafonds crédit')

    niveaux = [
        ('Standard', '50 – 300 USD', '6 mois max',  'Profil débutant, KYC de base'),
        ('Pro',      '50 – 700 USD', '9 mois max',  'Historique MM stable, KYC valide'),
        ('Pro+',     '50 – 1 200 USD','12 mois max', 'Bon comportement de remboursement'),
        ('Premium',  '50 – 2 500 USD','12 mois max', 'Profil solide, revenus vérifiés élevés'),
    ]

    tbl3 = doc.add_table(rows=1 + len(niveaux), cols=4)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(['Niveau', 'Plafond crédit', 'Durée max', 'Conditions']):
        cell = tbl3.cell(0, j)
        set_cell_bg(cell, '003C7A')
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = BLANC
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)

    niveau_colors = ['F4F4F4', 'E8F4EE', 'E8F0FA', 'FFF8E8']
    for i, (niveau, plafond, duree, cond) in enumerate(niveaux):
        row = tbl3.rows[i + 1]
        for j, txt in enumerate([niveau, plafond, duree, cond]):
            cell = row.cells[j]
            set_cell_bg(cell, niveau_colors[i])
            r = cell.paragraphs[0].add_run(txt)
            r.font.size = Pt(10)
            if j == 0:
                r.bold = True
                r.font.color.rgb = BLEU_RAWBANK
            elif j == 1:
                r.bold = True
                r.font.color.rgb = VERT_SCORE
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def section_exemple(doc):
    heading1(doc, '7. Exemple de présentation — Cas réel simulé')

    body(doc,
        'Voici comment Simbisa traite une demande de crédit, de bout en bout. '
        'Ce cas utilise Jean Kabila Mutombo, l\'un des profils de démonstration.')

    # Fiche client
    heading2(doc, 'Fiche du client')

    fiche = [
        ('Nom complet',       'Jean Kabila Mutombo'),
        ('Téléphone',         '+243 810 000 001'),
        ('Commune',           'Gombe, Kinshasa'),
        ('Profession',        'Commerçant'),
        ('Niveau de compte',  'Pro'),
        ('KYC',               'Validé — Carte électeur'),
        ('Revenu estimé',     '500 USD / mois'),
        ('Opérateur MM',      'Vodacom M-Pesa (081...)'),
        ('Âge',               '35 ans'),
    ]

    tbl = doc.add_table(rows=len(fiche), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(fiche):
        row = tbl.rows[i]
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.1)
        bg = 'F0F4FA' if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[0], 'E8F0FA')
        set_cell_bg(row.cells[1], bg)
        r0 = row.cells[0].paragraphs[0].add_run(k)
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = BLEU_RAWBANK
        r1 = row.cells[1].paragraphs[0].add_run(v)
        r1.font.size = Pt(10)
        row.cells[0].paragraphs[0].paragraph_format.space_after = Pt(3)
        row.cells[1].paragraphs[0].paragraph_format.space_after = Pt(3)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Demande
    heading2(doc, 'Sa demande de crédit')
    info_box(doc,
        'Demande soumise depuis l\'application mobile',
        'Montant : 250 USD   |   Durée : 6 mois   |   Motif : Achat stock boutique',
        'E8F4EE', VERT_SCORE)

    # Résultat du scoring
    heading2(doc, 'Résultat du scoring (calculé en 8 secondes)')

    score_details = [
        ('Moteur 1 — Règles & KYC',  22, 25, '1A6B3A'),
        ('Moteur 2 — Mobile Money',   24, 30, '007AC1'),
        ('Moteur 3 — Comportemental', 16, 20, 'E07A00'),
        ('Moteur 4 — IA / XGBoost',   20, 25, '550088'),
    ]

    for label, val, max_v, color in score_details:
        score_bar_table(doc, label, val, max_v, color)

    # Score final
    p_total = doc.add_paragraph()
    p_total.paragraph_format.space_before = Pt(8)
    p_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_total = p_total.add_run('SCORE FINAL : 82 / 100  →  APPROUVÉ')
    r_total.bold = True
    r_total.font.size = Pt(14)
    r_total.font.color.rgb = VERT_SCORE

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Mémo IA
    heading2(doc, 'Mémo généré par l\'IA (extrait)')
    info_box(doc,
        'Explication IA — Jean Kabila Mutombo, demande #1',
        'Le dossier de M. Jean Kabila Mutombo obtient un score de 82/100. '
        'Son compte M-Pesa enregistre 23 transactions sur 90 jours, avec des entrées '
        'moyennes de 312 000 CDF/mois et un ratio entrées/sorties de 1,38 — indicateur '
        'd\'une gestion budgétaire saine. Le KYC est validé, l\'âge (35 ans) est dans '
        'la plage autorisée, et le montant demandé (250 USD) est bien en deçà du plafond '
        'Pro (700 USD). Le modèle IA attribue une probabilité de défaut de 14 %, '
        'classant ce profil en risque FAIBLE. Variable la plus déterminante : '
        'régularité des entrées MM (SHAP +0.31). Recommandation : APPROBATION.',
        'EEF4FF')

    # Décision finale
    heading2(doc, 'Décision et suite du processus')

    etapes = [
        ('Décision automatique',   'Approuvé — score ≥ 70 (aucune intervention humaine nécessaire)'),
        ('Notification client',    'Jean reçoit une notification push : "Crédit de 250 USD approuvé"'),
        ('Déblocage',              'Le montant est crédité sur son wallet USD dans les 5 minutes'),
        ('Échéancier',             '6 mensualités de 44,17 USD (principal + 3 % intérêts/mois)'),
        ('Remboursement',          'Jean rembourse via M-Pesa ou directement depuis l\'application'),
        ('Clôture',                'Après le dernier paiement, le crédit est marqué "conclu". '
                                   'Son historique positif améliore son score pour la prochaine demande.'),
    ]

    for etape, desc in etapes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(f'{etape} : ')
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = BLEU_RAWBANK
        r2 = p.add_run(desc)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = GRIS_TEXTE


def section_impact(doc):
    heading1(doc, '8. Impact attendu et indicateurs clés')

    impacts = [
        ('Délai de traitement',
         'De 3-5 jours → 8 secondes pour les dossiers éligibles au traitement automatique.'),
        ('Taux d\'inclusion',
         'Des clients sans historique bancaire mais avec un compte Mobile Money actif '
         'deviennent éligibles au crédit.'),
        ('Réduction des impayés',
         'L\'analyse comportementale et l\'IA réduisent les approbations à risque '
         'élevé grâce à une évaluation objective et reproductible.'),
        ('Conformité réglementaire',
         'Chaque décision est tracée, horodatée et consultable par l\'auditeur. '
         'Le journal complet est disponible pour les inspections BCC.'),
        ('Économies opérationnelles',
         'Les agents se concentrent sur les dossiers complexes et les exceptions, '
         'réduisant le temps passé sur les demandes simples de 70 %.'),
    ]

    for titre, desc in impacts:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(f'  {titre}')
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = BLEU_RAWBANK
        p2 = doc.add_paragraph(f'  {desc}')
        p2.paragraph_format.left_indent  = Cm(0.5)
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after  = Pt(6)
        for run in p2.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = GRIS_TEXTE

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    heading2(doc, 'Canaux d\'accès')
    canaux = [
        ('Application mobile', 'Android & iOS', 'Interface complète : score, demandes, épargne, remboursements'),
        ('Tableau de bord web', 'Navigateur', 'Réservé aux équipes Rawbank (agents, managers, auditeurs)'),
        ('USSD *123#',  'Tout téléphone (sans internet)', 'Consulter solde, demander un crédit, voir son score'),
    ]

    tbl = doc.add_table(rows=1 + len(canaux), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['Canal', 'Support', 'Fonctions disponibles']):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, '003C7A')
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = BLANC
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)

    colors = ['E8F0FA', 'EBF5EE', 'FFF8EE']
    for i, (canal, support, fonctions) in enumerate(canaux):
        row = tbl.rows[i + 1]
        for j, txt in enumerate([canal, support, fonctions]):
            cell = row.cells[j]
            set_cell_bg(cell, colors[i])
            r = cell.paragraphs[0].add_run(txt)
            r.font.size = Pt(10)
            if j == 0:
                r.bold = True
                r.font.color.rgb = BLEU_RAWBANK
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)


def section_glossaire(doc):
    heading1(doc, '9. Glossaire simplifié')

    termes = [
        ('KYC', '"Know Your Customer" — vérification de l\'identité du client par un document officiel.'),
        ('Score de crédit', 'Note de 0 à 100 qui mesure la capacité estimée à rembourser un crédit.'),
        ('Mobile Money', 'Service financier via téléphone mobile : M-Pesa, Orange Money, Airtel Money, Africell.'),
        ('XGBoost', 'Algorithme d\'intelligence artificielle utilisé pour calculer la probabilité de défaut.'),
        ('SHAP', 'Technique qui explique quelle variable a le plus influencé la décision du modèle IA.'),
        ('RAG', 'Technologie qui permet à l\'IA de générer une explication en langage naturel basée sur les règles Rawbank.'),
        ('USSD', 'Canal de communication via réseau téléphonique (code *123#), fonctionne sans internet.'),
        ('Wallet', 'Portefeuille numérique en USD ou CDF sur la plateforme Simbisa.'),
        ('BCC', 'Banque Centrale du Congo — régulateur bancaire de la RDC.'),
        ('Taux de défaut', 'Pourcentage de clients qui ne remboursent pas leur crédit dans les délais.'),
        ('Exception crédit', 'Demande hors normes (ex. : montant supérieur au plafond automatique) '
                             'nécessitant une validation manuelle du responsable crédit.'),
    ]

    tbl = doc.add_table(rows=len(termes), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (terme, definition) in enumerate(termes):
        row = tbl.rows[i]
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.8)
        bg = 'F0F4FA' if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[0], 'E8F0FA')
        set_cell_bg(row.cells[1], bg)
        r0 = row.cells[0].paragraphs[0].add_run(terme)
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = BLEU_RAWBANK
        r1 = row.cells[1].paragraphs[0].add_run(definition)
        r1.font.size = Pt(10)
        r1.font.color.rgb = GRIS_TEXTE
        row.cells[0].paragraphs[0].paragraph_format.space_after = Pt(3)
        row.cells[1].paragraphs[0].paragraph_format.space_after = Pt(3)


# ─── Main ─────────────────────────────────────────────────────────────────────
def main():
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Police par défaut
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    cover_page(doc)
    section_contexte(doc)
    doc.add_page_break()
    section_problematique(doc)
    section_situation(doc)
    doc.add_page_break()
    section_utilisateurs(doc)
    doc.add_page_break()
    section_score(doc)
    doc.add_page_break()
    section_identifiants(doc)
    doc.add_page_break()
    section_exemple(doc)
    doc.add_page_break()
    section_impact(doc)
    section_glossaire(doc)

    out = 'Simbisa_Presentation_NonTechnique.docx'
    doc.save(out)
    print(f'Document généré : {out}')


if __name__ == '__main__':
    main()
